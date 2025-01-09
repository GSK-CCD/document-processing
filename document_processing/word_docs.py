from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml.table import CT_Tbl as table_type
from docx.oxml.text.paragraph import CT_P as paragraph_type
from docx.table import Table
from starlette import datastructures

from document_processing.base import BaseFileProcessor


class WordDocProcessor(BaseFileProcessor):

    async def _get_doc(self):
        if isinstance(self.file_name, str):
            return Document(self.file_name)
        elif isinstance(self.file_name, datastructures.UploadFile):
            page = await self.file_name.read()
            doc = Document(BytesIO(page))
            return doc
        else:
            raise ValueError(f"file_name must be either a string or a datastructures.UploadFile, got '{type(self.file_name)}'")

    async def extract_text(self, target_column: Optional[int] = None) -> str:
        """
        Takes a Word file and extracts the text from it to a string.
        """
        doc = await self._get_doc()
        full_text = ""
        tables = doc.tables
        for element in doc.element.body:
            if isinstance(element, paragraph_type):
                full_text += "\n" + element.text
            elif isinstance(element, table_type):
                full_text += "\n" + self.extract_from_table(tables.pop(0), target_column)
        return full_text

    def extract_from_table(self, table: Table, target_column: Optional[int]) -> str:
        if target_column is not None:
            return self._extract_specific_column(table, target_column)
        return self._table_to_markdown(table)

    def _extract_specific_column(self, table: Table, target_column: int) -> str:
        table_text = ""
        for row in table.rows:
            if len(row.cells) == 1:
                table_cell = row.cells[0]
            else:
                table_cell = row.cells[target_column]
            table_text += table_cell.text
        return table_text

    def _table_to_markdown(self, table: Table) -> str:
        table_text = ""
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                table_text += f"| {cell.text} "
                if j == len(row.cells) - 1:
                    table_text += "|\n"
            if i == 0:
                table_text += "| --- " * len(row.cells) + "|\n"
        return table_text
